import io
import re
import streamlit as st
import pdfplumber
from docx import Document

# Optional: Google Gemini (only needed when you click "Tailor Resume")
try:
    import google.generativeai as genai
except Exception:
    genai = None

st.set_page_config(page_title="ATS Resume Tailor", page_icon="🧠", layout="centered")

st.title("🧠 ATS Resume Tailor")
st.write("Upload your resume, paste a job description, and get a tailored, ATS-friendly version.")

# --- Helpers ---
def read_docx(file_like) -> str:
    doc = Document(file_like)
    return "\n".join(p.text for p in doc.paragraphs)

def read_pdf(file_like) -> str:
    text = []
    with pdfplumber.open(file_like) as pdf:
        for page in pdf.pages:
            text.append(page.extract_text() or "")
    return "\n".join(text).strip()

def extract_resume_text(upload):
    name = (upload.name or "").lower()
    if name.endswith(".docx"):
        return read_docx(upload)
    if name.endswith(".pdf"):
        return read_pdf(upload)
    # Fallback: treat as plain text
    return upload.read().decode("utf-8", errors="ignore")

def save_docx(text: str) -> bytes:
    buf = io.BytesIO()
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(buf)
    buf.seek(0)
    return buf.read()

def ats_lint(text: str):
    """Light checks for ATS-friendliness."""
    warnings = []
    # Fancy columns/tables impossible to detect here, but we can flag long lines & odd bullets
    long_lines = [ln for ln in text.splitlines() if len(ln) > 160]
    if long_lines:
        warnings.append(f"{len(long_lines)} lines exceed ~160 chars; consider concise bullets.")
    if re.search(r"[^\x00-\x7F]", text):
        warnings.append("Non-ASCII characters found; stick to plain text where possible.")
    if not re.search(r"(?i)skills|experience|education", text):
        warnings.append("Missing standard headings like Skills, Experience, Education.")
    return warnings

# --- UI ---
with st.sidebar:
    st.header("Settings")
    provider = st.selectbox("Model", ["Gemini 1.5 Flash"], index=0)
    api_key = st.text_input("GEMINI_API_KEY (or set via Streamlit Secrets)", type="password")
    use_secrets = st.checkbox("Use Streamlit Secrets", True)
    if use_secrets:
        api_key = st.secrets.get("GEMINI_API_KEY", api_key)

resume_file = st.file_uploader("Upload Resume (PDF or DOCX preferred)", type=["pdf", "docx", "txt"], accept_multiple_files=False)
jd_text = st.text_area("Paste Job Description", height=220, placeholder="Paste the JD here…")

# Optional manual paste of resume text
with st.expander("Or paste resume text instead"):
    resume_text_manual = st.text_area("Resume (plain text)", height=220)

col_a, col_b = st.columns(2)
with col_a:
    draft_btn = st.button("✨ Tailor Resume")
with col_b:
    clear_btn = st.button("Clear")

if clear_btn:
    st.experimental_rerun()

# --- Main flow ---
if draft_btn:
    # 1) Get resume text
    resume_text = ""
    if resume_file is not None:
        try:
            resume_text = extract_resume_text(resume_file)
        except Exception as e:
            st.error(f"Could not read resume: {e}")
    if not resume_text and resume_text_manual.strip():
        resume_text = resume_text_manual.strip()

    if not resume_text:
        st.warning("Please upload a resume or paste your resume text.")
        st.stop()
    if not jd_text.strip():
        st.warning("Please paste the job description.")
        st.stop()

    if genai is None:
        st.error("google-generativeai not installed.")
        st.stop()
    if not api_key:
        st.info("Add your GEMINI_API_KEY in the sidebar or Streamlit Secrets.")
        st.stop()

    # 2) Configure model
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-1.5-flash")

    # 3) Prompt — strict ATS formatting + honesty
    system_rules = """
You are an expert resume writer specializing in ATS compliance.

Rules:
- Keep it factual. Do NOT invent experience, employers, or dates.
- Use plain, ATS-safe formatting: no tables, columns, images, headers/footers.
- Standard section headings: SUMMARY, SKILLS, EXPERIENCE, EDUCATION, CERTIFICATIONS (if any), PROJECTS (optional).
- Bullets should be concise and results-oriented. Prefer '-' bullets.
- Mirror relevant keywords/phrases from the JD naturally (no stuffing).
- Keep to 1–2 pages of text.
- Preserve job titles/employers/dates from the original resume unless user content provides updates.
- Where there is a relevant achievement you can rephrase it using JD language, but do not fabricate.

Return TWO blocks in this exact format:

<RESUME>
[ATS-optimized resume only]
</RESUME>

<MATCH_REPORT>
- Top keywords used
- Notable gaps (if any)
- Suggestions to strengthen fit (skills, certs, metrics)
</MATCH_REPORT>
""".strip()

    user_payload = f"""
[RESUME]
{resume_text}

[JOB_DESCRIPTION]
{jd_text}
""".strip()

    with st.spinner("Tailoring your resume…"):
        resp = model.generate_content([system_rules, user_payload])
        content = resp.text or ""

    # 4) Parse response
    resume_block = ""
    report_block = ""
    m1 = re.search(r"<RESUME>(.*?)</RESUME>", content, flags=re.DOTALL | re.IGNORECASE)
    m2 = re.search(r"<MATCH_REPORT>(.*?)</MATCH_REPORT>", content, flags=re.DOTALL | re.IGNORECASE)
    if m1:
        resume_block = m1.group(1).strip()
    if m2:
        report_block = m2.group(1).strip()
    if not resume_block:
        resume_block = content.strip()

    # 5) ATS lint
    warnings = ats_lint(resume_block)
    if warnings:
        st.warning("ATS Lint Suggestions:\n- " + "\n- ".join(warnings))

    # 6) Show results
    st.subheader("Tailored Resume (ATS-friendly)")
    st.code(resume_block, language="text")

    st.subheader("Match Report")
    st.write(report_block if report_block else "—")

    # 7) Downloads
    st.download_button(
        "Download as .txt",
        resume_block.encode("utf-8"),
        file_name="ATS_Resume.txt",
        mime="text/plain",
    )

    docx_bytes = save_docx(resume_block)
    st.download_button(
        "Download as .docx",
        docx_bytes,
        file_name="ATS_Resume.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

st.caption("Tip: keep each bullet to one line and include measurable impact (%, time saved, cost reduced).")
